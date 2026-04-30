"""
Document Generator — creates .docx files from campaign results.

Generates a professionally formatted Word document containing:
- Campaign header (name, date, status, persona info)
- Configuration section (formatted with CampaignParser)
- Tweets section with analysis scores and generated comments
- Legacy analysis section (for backward compatibility)

Property 20 from design.md: the document must contain all required sections.
"""

from __future__ import annotations

import os
import tempfile
from datetime import datetime, timezone
from typing import List, Optional, Dict

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.models.campaign import Campaign, Tweet
from src.models.analysis import Analysis
from src.models.tweet_analysis import TweetAnalysis
from src.models.tweet_comment import TweetComment
from src.models.communication_style import CommunicationStyle
from src.services.campaign_parser import CampaignParser
from src.core.logging_config import get_logger

logger = get_logger("services.document_generator")


class DocumentGenerator:
    """Generates .docx campaign result documents using python-docx."""

    def generate(
        self,
        campaign: Campaign,
        tweets: List[Tweet],
        analysis: Analysis,
        tweet_analyses: Optional[List[TweetAnalysis]] = None,
        tweet_comments: Optional[List[TweetComment]] = None,
        persona: Optional[CommunicationStyle] = None,
    ) -> str:
        """
        Create a .docx file and return its temporary file path.

        When tweet_analyses and tweet_comments are provided, the document
        includes per-tweet scores and generated comments. Falls back to the
        legacy analysis section when they are absent.

        The caller is responsible for cleaning up the file after upload.
        """
        doc = Document()
        self._add_header(doc, campaign, persona)
        self._add_configuration_section(doc, campaign)

        if tweet_analyses or tweet_comments:
            self._add_enriched_tweets_section(doc, tweets, tweet_analyses, tweet_comments)
        else:
            self._add_tweets_section(doc, tweets)
            if analysis:
                self._add_analysis_section(doc, analysis)

        # Save to a named temp file so it can be read back for upload
        tmp = tempfile.NamedTemporaryFile(
            suffix=".docx",
            prefix=f"campaign_{campaign.id}_",
            delete=False,
        )
        doc.save(tmp.name)
        tmp.close()
        logger.info("Document generated at %s (%d tweets)", tmp.name, len(tweets))
        return tmp.name

    # ─── Sections ────────────────────────────────────────────────────────────

    def _add_header(
        self, doc: Document, campaign: Campaign, persona: Optional[CommunicationStyle] = None
    ) -> None:
        """Add campaign title, date, status, and optional communication style info."""
        title = doc.add_heading(campaign.name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(
            f"Generated: {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  |  "
            f"Status: {campaign.status.value.upper()}  |  "
            f"Results: {len(campaign.config.profiles or campaign.config.keywords or [])} sources"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if persona:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"Estilo de Comunicação: {persona.name} — {persona.title}")
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

        doc.add_paragraph()  # spacer

    def _add_configuration_section(self, doc: Document, campaign: Campaign) -> None:
        """Add the 'Configuration Used' section."""
        doc.add_heading("Configuration Used", level=2)
        cfg = campaign.config

        rows = [
            ("Search Type", campaign.search_type),
            ("Language", cfg.language),
        ]

        if cfg.profiles:
            rows.append(("Profiles", CampaignParser.format_profiles(cfg.profiles)))
        if cfg.keywords:
            rows.append(("Keywords", CampaignParser.format_keywords(cfg.keywords)))

        rows.append(
            (
                "Engagement Filters",
                CampaignParser.format_engagement_filters(
                    cfg.min_likes, cfg.min_retweets, cfg.min_replies
                ),
            )
        )
        rows.append(("Time Window", f"Last {cfg.days_back} day{'s' if cfg.days_back > 1 else ''}"))

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(rows):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            table.cell(i, 0).paragraphs[0].runs[0].bold = True

        doc.add_paragraph()  # spacer

    def _add_enriched_tweets_section(
        self,
        doc: Document,
        tweets: List[Tweet],
        tweet_analyses: Optional[List[TweetAnalysis]],
        tweet_comments: Optional[List[TweetComment]],
    ) -> None:
        """Add tweets with per-tweet analysis scores and generated comments."""
        # Build lookup maps keyed by tweet_id
        analysis_map: Dict[str, TweetAnalysis] = {}
        if tweet_analyses:
            for a in tweet_analyses:
                analysis_map[a.tweet_id] = a

        comment_map: Dict[str, TweetComment] = {}
        if tweet_comments:
            for c in tweet_comments:
                comment_map[c.tweet_id] = c

        # Separate top-3 from the rest
        top3_ids = {tid for tid, a in analysis_map.items() if a.is_top_3}

        doc.add_heading(f"Collected Tweets ({len(tweets)})", level=2)

        if not tweets:
            doc.add_paragraph("No tweets were collected for this campaign.")
            return

        for i, tweet in enumerate(tweets, 1):
            tweet_id = tweet.id
            analysis = analysis_map.get(tweet_id)
            comment = comment_map.get(tweet_id)
            is_top3 = tweet_id in top3_ids

            # Tweet header
            p = doc.add_paragraph()
            header_text = f"{i}. @{tweet.author}"
            if is_top3:
                header_text = f"⭐ {header_text} [TOP 3]"
            p.add_run(header_text).bold = True
            p.add_run(
                f"  ❤️ {tweet.likes}  🔁 {tweet.reposts}  💬 {tweet.replies}"
                f"  —  {tweet.timestamp.strftime('%Y-%m-%d %H:%M UTC')}"
            ).font.size = Pt(9)

            # Tweet text
            doc.add_paragraph(tweet.text)

            # Tweet URL
            link_p = doc.add_paragraph()
            link_p.add_run(tweet.url).font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

            # Analysis scores
            if analysis:
                scores_p = doc.add_paragraph()
                scores_p.add_run("Analysis: ").bold = True
                scores_p.add_run(
                    f"Score {analysis.average_score:.1f}/10  |  "
                    f"Verdict: {analysis.verdict.value}  |  "
                    f"Relevance: {analysis.lead_relevance_score}  "
                    f"Tone: {analysis.tone_of_voice_score}  "
                    f"Insight: {analysis.insight_strength_score}  "
                    f"Engagement: {analysis.engagement_potential_score}  "
                    f"Safety: {analysis.brand_safety_score}"
                ).font.size = Pt(9)
                if analysis.notes:
                    note_p = doc.add_paragraph()
                    note_p.add_run("Notes: ").bold = True
                    note_p.add_run(analysis.notes).font.size = Pt(9)

            # Generated comment
            if comment:
                comment_p = doc.add_paragraph()
                comment_p.add_run("Generated Comment: ").bold = True
                comment_p.add_run(comment.comment_text)
                char_p = doc.add_paragraph()
                char_p.add_run(f"({comment.char_count} characters)").font.size = Pt(8)
                char_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            if i < len(tweets):
                doc.add_paragraph("─" * 60)

        doc.add_paragraph()  # spacer

    def _add_tweets_section(self, doc: Document, tweets: List[Tweet]) -> None:
        """Add the collected tweets section (legacy, no analysis/comments)."""
        doc.add_heading(f"Collected Tweets ({len(tweets)})", level=2)

        if not tweets:
            doc.add_paragraph("No tweets were collected for this campaign.")
            return

        for i, tweet in enumerate(tweets, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. @{tweet.author}").bold = True
            p.add_run(
                f"  ❤️ {tweet.likes}  🔁 {tweet.reposts}  💬 {tweet.replies}"
                f"  —  {tweet.timestamp.strftime('%Y-%m-%d %H:%M UTC')}"
            ).font.size = Pt(9)

            doc.add_paragraph(tweet.text)

            link_p = doc.add_paragraph()
            link_p.add_run(tweet.url).font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            link_p.add_run("").font.size = Pt(8)

            if i < len(tweets):
                doc.add_paragraph("─" * 60)

        doc.add_paragraph()  # spacer

    def _add_analysis_section(self, doc: Document, analysis: Optional[Analysis]) -> None:
        """Add the OpenAI analysis section (legacy)."""
        if not analysis:
            return
        doc.add_heading("Analysis", level=2)

        doc.add_heading("Summary", level=3)
        doc.add_paragraph(analysis.summary)

        if analysis.key_themes:
            doc.add_heading("Key Themes", level=3)
            for theme in analysis.key_themes:
                doc.add_paragraph(theme, style="List Bullet")

        doc.add_heading("Sentiment", level=3)
        doc.add_paragraph(analysis.sentiment.capitalize())

        if analysis.top_influencers:
            doc.add_heading("Top Influencers", level=3)
            for influencer in analysis.top_influencers:
                doc.add_paragraph(influencer, style="List Bullet")

        if analysis.recommendations:
            doc.add_heading("Recommendations", level=3)
            for rec in analysis.recommendations:
                doc.add_paragraph(rec, style="List Bullet")
