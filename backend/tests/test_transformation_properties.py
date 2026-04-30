"""
Property-Based Tests for Transformation Properties

Tests the correctness properties 18-20 defined in design.md using Hypothesis.
These tests verify that the TwitterScrapingEngine's filtering and transformation
methods, and DocumentGenerator, maintain their invariants across a wide range of inputs.
"""

# Set up environment variables BEFORE any imports that might load settings
import os

os.environ.setdefault("SUPABASE_URL", "https://test.supabase.co")
os.environ.setdefault("SUPABASE_KEY", "test-key")
os.environ.setdefault("REDIS_URL", "redis://localhost:6379/1")
os.environ.setdefault("CELERY_BROKER_URL", "redis://localhost:6379/1")
os.environ.setdefault("CELERY_RESULT_BACKEND", "redis://localhost:6379/1")
os.environ.setdefault("ENCRYPTION_KEY", "test-encryption-key-32-bytes-long!!")
os.environ.setdefault("DEBUG", "True")

from hypothesis import given, strategies as st, settings
from datetime import datetime, timezone
from unittest.mock import Mock
from uuid import uuid4
import tempfile
import pytest

# Now safe to import after environment is set up
from src.services.scraping_engine import TwitterScrapingEngine
from src.services.document_generator import DocumentGenerator
from src.models.campaign import (
    ScrapingConfig,
    SearchType,
    Tweet,
    Campaign,
    CampaignConfig,
    CampaignStatus,
)
from src.models.analysis import Analysis


# ─── Hypothesis Strategies ───────────────────────────────────────────────────


# Strategy for generating engagement counts
engagement_strategy = st.integers(min_value=0, max_value=10000)


# Strategy for generating XML-safe text (no control characters)
def xml_safe_text(min_size=1, max_size=100):
    """Generate XML-safe text without control characters"""
    return st.text(
        alphabet=st.characters(
            blacklist_categories=("Cc", "Cs", "Cn"),  # Control, surrogate, not assigned
            blacklist_characters="\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f",
        ),
        min_size=min_size,
        max_size=max_size,
    ).filter(lambda x: x.strip() if min_size > 0 else True)


# Strategy for generating tweet IDs
tweet_id_strategy = st.text(
    alphabet=st.characters(whitelist_categories=("Nd",)), min_size=10, max_size=20
)

# Strategy for generating Twitter usernames
username_strategy = st.text(
    alphabet=st.characters(whitelist_categories=("Lu", "Ll", "Nd"), whitelist_characters="_"),
    min_size=1,
    max_size=20,
)

# Strategy for generating tweet text (avoid control characters for XML compatibility)
tweet_text_strategy = xml_safe_text(min_size=1, max_size=280)

# Strategy for generating timestamps
timestamp_strategy = st.datetimes(
    min_value=datetime(2020, 1, 1), max_value=datetime(2024, 12, 31)
).map(lambda dt: dt.replace(tzinfo=timezone.utc))


# Strategy for generating raw tweet data from Apify
@st.composite
def raw_tweet_strategy(draw):
    """Generate raw tweet data as returned by Apify"""
    return {
        "id": draw(tweet_id_strategy),
        "url": f"https://twitter.com/user/status/{draw(tweet_id_strategy)}",
        "author": {"userName": draw(username_strategy)},
        "text": draw(tweet_text_strategy),
        "likeCount": draw(engagement_strategy),
        "retweetCount": draw(engagement_strategy),
        "replyCount": draw(engagement_strategy),
        "createdAt": draw(timestamp_strategy).isoformat(),
    }


# Strategy for generating scraping configs
@st.composite
def scraping_config_strategy(draw):
    """Generate a valid scraping configuration"""
    return ScrapingConfig(
        search_type=SearchType.KEYWORDS,
        keywords=["test"],
        language="en",
        min_likes=draw(engagement_strategy),
        min_retweets=draw(engagement_strategy),
        min_replies=draw(engagement_strategy),
        hours_back=24,
        apify_token="test-token",
    )


# Strategy for generating Tweet models
@st.composite
def tweet_model_strategy(draw):
    """Generate a valid Tweet model"""
    return Tweet(
        id=draw(tweet_id_strategy),
        url=f"https://twitter.com/user/status/{draw(tweet_id_strategy)}",
        author=draw(username_strategy),
        text=draw(tweet_text_strategy),
        likes=draw(engagement_strategy),
        reposts=draw(engagement_strategy),
        replies=draw(engagement_strategy),
        timestamp=draw(timestamp_strategy),
    )


# Strategy for generating Campaign models
@st.composite
def campaign_strategy(draw):
    """Generate a valid Campaign model"""
    search_type = draw(st.sampled_from([SearchType.PROFILE, SearchType.KEYWORDS]))
    return Campaign(
        id=uuid4(),
        name=draw(xml_safe_text(min_size=1, max_size=100)),
        social_network="twitter",
        search_type=search_type.value,
        config=CampaignConfig(
            profiles=["user1", "user2"] if search_type == SearchType.PROFILE else None,
            keywords=["keyword1", "keyword2"] if search_type == SearchType.KEYWORDS else None,
            language="en",
            min_likes=draw(engagement_strategy),
            min_retweets=draw(engagement_strategy),
            min_replies=draw(engagement_strategy),
            hours_back=24,
        ),
        status=CampaignStatus.COMPLETED,
        error_message=None,
        document_url=None,
        results_count=0,
        created_at=datetime.now(tz=timezone.utc),
        updated_at=datetime.now(tz=timezone.utc),
        completed_at=datetime.now(tz=timezone.utc),
    )


# Strategy for generating Analysis models
@st.composite
def analysis_strategy(draw):
    """Generate a valid Analysis model"""
    return Analysis(
        summary=draw(xml_safe_text(min_size=10, max_size=500)),
        key_themes=draw(st.lists(xml_safe_text(min_size=1, max_size=50), min_size=1, max_size=5)),
        sentiment=draw(st.sampled_from(["positive", "negative", "neutral", "mixed"])),
        top_influencers=draw(st.lists(username_strategy, min_size=0, max_size=5)),
        recommendations=draw(
            st.lists(xml_safe_text(min_size=10, max_size=200), min_size=1, max_size=5)
        ),
    )


# ─── Property 18: Local Filtering Enforces Engagement Criteria ──────────────


@given(
    raw_tweets=st.lists(raw_tweet_strategy(), min_size=1, max_size=50),
    config=scraping_config_strategy(),
)
@settings(max_examples=20, deadline=None)
def test_property_18_local_filtering_enforces_engagement_criteria(raw_tweets, config):
    """
    **Validates: Requirements 5.11**

    Property 18: Local Filtering Enforces Engagement Criteria

    For any list of tweets and engagement criteria (min_likes, min_retweets,
    min_replies), filtering SHALL remove all tweets that do not meet ALL
    specified criteria.
    """
    # Create engine with mock client
    mock_client = Mock()
    engine = TwitterScrapingEngine(mock_client)

    # Apply filters
    filtered = engine.apply_filters(raw_tweets, config)

    # Verify all filtered tweets meet ALL criteria
    for tweet in filtered:
        likes = tweet.get("likeCount", tweet.get("likes", 0)) or 0
        retweets = tweet.get("retweetCount", tweet.get("retweets", 0)) or 0
        replies = tweet.get("replyCount", tweet.get("replies", 0)) or 0

        assert (
            likes >= config.min_likes
        ), f"Filtered tweet has {likes} likes, but min_likes is {config.min_likes}"
        assert (
            retweets >= config.min_retweets
        ), f"Filtered tweet has {retweets} retweets, but min_retweets is {config.min_retweets}"
        assert (
            replies >= config.min_replies
        ), f"Filtered tweet has {replies} replies, but min_replies is {config.min_replies}"

    # Verify that tweets not meeting criteria are excluded
    for tweet in raw_tweets:
        likes = tweet.get("likeCount", tweet.get("likes", 0)) or 0
        retweets = tweet.get("retweetCount", tweet.get("retweets", 0)) or 0
        replies = tweet.get("replyCount", tweet.get("replies", 0)) or 0

        meets_criteria = (
            likes >= config.min_likes
            and retweets >= config.min_retweets
            and replies >= config.min_replies
        )

        if meets_criteria:
            # Tweet should be in filtered list
            assert (
                tweet in filtered
            ), "Tweet meeting all criteria should be included in filtered results"
        else:
            # Tweet should NOT be in filtered list
            assert (
                tweet not in filtered
            ), "Tweet not meeting all criteria should be excluded from filtered results"


# ─── Property 19: Tweet Transformation Includes All Required Fields ─────────


@given(raw_tweets=st.lists(raw_tweet_strategy(), min_size=1, max_size=50))
@settings(max_examples=20, deadline=None)
def test_property_19_tweet_transformation_includes_all_required_fields(raw_tweets):
    """
    **Validates: Requirements 5.12**

    Property 19: Tweet Transformation Includes All Required Fields

    For any raw tweet data from Apify, the transformed tweet SHALL contain
    all required fields: id, url, author, text, likes, reposts, replies,
    timestamp.
    """
    # Create engine with mock client
    mock_client = Mock()
    engine = TwitterScrapingEngine(mock_client)

    # Transform tweets
    transformed = engine.transform_results(raw_tweets)

    # Verify all transformed tweets have all required fields
    required_fields = ["id", "url", "author", "text", "likes", "reposts", "replies", "timestamp"]

    for tweet in transformed:
        # Check that tweet is a Tweet model instance
        assert isinstance(tweet, Tweet), "Transformed tweet should be a Tweet model instance"

        # Check all required fields are present and not None
        for field in required_fields:
            assert hasattr(tweet, field), f"Transformed tweet missing required field: {field}"

            value = getattr(tweet, field)
            assert (
                value is not None
            ), f"Transformed tweet has None value for required field: {field}"

            # Additional type checks
            if field == "id":
                assert isinstance(value, str), "Tweet id should be a string"
                assert len(value) > 0, "Tweet id should not be empty"
            elif field == "url":
                assert isinstance(value, str), "Tweet url should be a string"
                assert len(value) > 0, "Tweet url should not be empty"
            elif field == "author":
                assert isinstance(value, str), "Tweet author should be a string"
                assert len(value) > 0, "Tweet author should not be empty"
            elif field == "text":
                assert isinstance(value, str), "Tweet text should be a string"
            elif field in ["likes", "reposts", "replies"]:
                assert isinstance(value, int), f"Tweet {field} should be an integer"
                assert value >= 0, f"Tweet {field} should be non-negative"
            elif field == "timestamp":
                assert isinstance(value, datetime), "Tweet timestamp should be a datetime"
                assert value.tzinfo is not None, "Tweet timestamp should be timezone-aware"


# ─── Property 20: Document Contains All Required Sections ───────────────────


@given(
    campaign=campaign_strategy(),
    tweets=st.lists(tweet_model_strategy(), min_size=1, max_size=20),
    analysis=analysis_strategy(),
)
@settings(max_examples=20, deadline=None)
def test_property_20_document_contains_all_required_sections(campaign, tweets, analysis):
    """
    **Validates: Requirements 7.3, 7.4**

    Property 20: Document Contains All Required Sections

    For any campaign with tweets and analysis, the generated document SHALL
    contain sections for: campaign name, execution date, configuration
    parameters, collected tweets, and analysis results.
    """
    # Create document generator
    generator = DocumentGenerator()

    # Generate document
    doc_path = generator.generate(campaign, tweets, analysis)

    try:
        # Verify file was created
        assert os.path.exists(doc_path), "Document file should be created"
        assert doc_path.endswith(".docx"), "Document should be a .docx file"

        # Read the document to verify sections
        from docx import Document

        doc = Document(doc_path)

        # Extract all text from the document (paragraphs + tables)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        # Verify campaign name is present (header section)
        assert campaign.name in full_text, "Document should contain campaign name"

        # Verify execution date is present (header section)
        assert "Generated:" in full_text, "Document should contain generation date"

        # Verify status is present (header section)
        assert "Status:" in full_text, "Document should contain campaign status"

        # Verify configuration section is present
        assert (
            "Configuration Used" in full_text
        ), "Document should contain 'Configuration Used' section"

        # Verify configuration parameters are present
        assert "Search Type" in full_text, "Document should contain search type"
        assert "Language" in full_text, "Document should contain language"
        assert (
            "Engagement Filters" in full_text or "min_likes" in full_text.lower()
        ), "Document should contain engagement filters"
        assert (
            "Time Window" in full_text or "hours" in full_text.lower()
        ), "Document should contain time window"

        # Verify tweets section is present
        assert "Collected Tweets" in full_text, "Document should contain 'Collected Tweets' section"

        # Verify at least some tweet content is present
        # (Check for at least one tweet author or text)
        tweet_content_found = any(tweet.author in full_text for tweet in tweets[:5])
        assert tweet_content_found, "Document should contain tweet content"

        # Verify analysis section is present
        assert "Analysis" in full_text, "Document should contain 'Analysis' section"

        # Verify analysis subsections are present
        assert "Summary" in full_text, "Document should contain analysis summary"
        assert "Sentiment" in full_text, "Document should contain sentiment analysis"

        # Verify analysis content is present
        assert (
            analysis.summary in full_text or analysis.sentiment in full_text
        ), "Document should contain analysis content"

    finally:
        # Clean up temporary file
        if os.path.exists(doc_path):
            os.unlink(doc_path)


# ─── Additional Property: Filtering Preserves Order ──────────────────────────


@given(
    raw_tweets=st.lists(raw_tweet_strategy(), min_size=2, max_size=20),
    config=scraping_config_strategy(),
)
@settings(max_examples=10, deadline=None)
def test_filtering_preserves_order(raw_tweets, config):
    """
    Additional property: Filtering should preserve the relative order of tweets.

    If tweet A appears before tweet B in the input, and both pass the filters,
    then tweet A should appear before tweet B in the output.
    """
    # Create engine with mock client
    mock_client = Mock()
    engine = TwitterScrapingEngine(mock_client)

    # Apply filters
    filtered = engine.apply_filters(raw_tweets, config)

    # Build a mapping of tweet to original index
    original_indices = {id(tweet): i for i, tweet in enumerate(raw_tweets)}

    # Verify order is preserved
    for i in range(len(filtered) - 1):
        current_idx = original_indices[id(filtered[i])]
        next_idx = original_indices[id(filtered[i + 1])]

        assert current_idx < next_idx, "Filtering should preserve the original order of tweets"


# ─── Additional Property: Transformation Is Idempotent ───────────────────────


@given(raw_tweets=st.lists(raw_tweet_strategy(), min_size=1, max_size=20))
@settings(max_examples=10, deadline=None)
def test_transformation_handles_malformed_data_gracefully(raw_tweets):
    """
    Additional property: Transformation should handle malformed data gracefully
    by skipping invalid tweets rather than crashing.
    """
    # Create engine with mock client
    mock_client = Mock()
    engine = TwitterScrapingEngine(mock_client)

    # Add some malformed tweets to the list
    malformed_tweets = raw_tweets + [
        {},  # Empty dict
        {"id": "123"},  # Missing required fields
        {"text": "test"},  # Missing id
    ]

    # Transform should not raise an exception
    try:
        transformed = engine.transform_results(malformed_tweets)

        # Transformed list should contain only valid tweets
        # (may be fewer than input if some were malformed)
        assert len(transformed) <= len(
            malformed_tweets
        ), "Transformation should not create more tweets than input"

        # All transformed tweets should be valid Tweet instances
        for tweet in transformed:
            assert isinstance(tweet, Tweet), "All transformed tweets should be Tweet instances"
            # ID can be empty string if the raw data had empty/invalid ID
            # but it should be a string
            assert isinstance(tweet.id, str), "Tweet id should be a string"
            assert isinstance(tweet.url, str), "Tweet url should be a string"
            assert isinstance(tweet.author, str), "Tweet author should be a string"

    except Exception as e:
        # If an exception is raised, fail the test
        pytest.fail(f"Transformation should handle malformed data gracefully, but raised: {e}")


# ─── Additional Property: Document Generation Is Deterministic ───────────────


@given(
    campaign=campaign_strategy(),
    tweets=st.lists(tweet_model_strategy(), min_size=1, max_size=10),
    analysis=analysis_strategy(),
)
@settings(max_examples=20, deadline=None)
def test_document_generation_creates_valid_docx(campaign, tweets, analysis):
    """
    Additional property: Document generation should always create a valid
    .docx file that can be opened and read.
    """
    # Create document generator
    generator = DocumentGenerator()

    # Generate document
    doc_path = generator.generate(campaign, tweets, analysis)

    try:
        # Verify file exists and has content
        assert os.path.exists(doc_path), "Document file should exist"
        assert os.path.getsize(doc_path) > 0, "Document file should not be empty"

        # Verify it's a valid .docx file by opening it
        from docx import Document

        doc = Document(doc_path)

        # Verify document has content
        assert len(doc.paragraphs) > 0, "Document should have paragraphs"

        # Verify document has some text
        text_content = "\n".join([para.text for para in doc.paragraphs])
        assert len(text_content.strip()) > 0, "Document should have text content"

    except Exception as e:
        pytest.fail(f"Generated document should be a valid .docx file, but got error: {e}")

    finally:
        # Clean up temporary file
        if os.path.exists(doc_path):
            os.unlink(doc_path)
